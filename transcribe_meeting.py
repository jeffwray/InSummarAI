import os
import json
from openai import OpenAI
from docx import Document
import argparse
import logging
from pyannote.audio import Pipeline
import torchaudio
import ffmpeg
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import io

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Initialize the OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Helper function to load intermediate results
def load_checkpoint(file_path):
    if os.path.exists(file_path):
        with open(file_path, 'r') as f:
            logging.info(f"Loaded existing checkpoint from {file_path}")
            return json.load(f)
    return None

# Helper function to save intermediate results
def save_checkpoint(data, file_path):
    with open(file_path, 'w') as f:
        json.dump(data, f)
        logging.info(f"Checkpoint saved to {file_path}")

# Transcribing audio with Whisper API
def transcribe_audio(audio_file_path):
    transcription_file = f"{audio_file_path}_transcription.json"
    
    # Check if transcription already exists
    transcription = load_checkpoint(transcription_file)
    if transcription:
        return transcription
    
    logging.info(f"Starting transcription for {audio_file_path}")
    try:
        with open(audio_file_path, 'rb') as audio_file:
            response = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file,
                response_format="text"
            )
        # The response is now directly the transcription text
        transcription = response
        save_checkpoint(transcription, transcription_file)
        logging.info("Transcription completed successfully.")
        return transcription
    except Exception as e:
        logging.error(f"Error during transcription: {e}")
        raise e

# Perform speaker diarization
def perform_speaker_diarization(audio_file_path):
    try:
        logging.info("Starting speaker diarization for {}".format(audio_file_path))
        
        # Convert audio to WAV format if it's not already
        wav_file_path = audio_file_path
        if not audio_file_path.lower().endswith('.wav'):
            wav_file_path = audio_file_path.rsplit('.', 1)[0] + '.wav'
            logging.info(f"Converting {audio_file_path} to WAV format: {wav_file_path}")
            (
                ffmpeg
                .input(audio_file_path)
                .output(wav_file_path, acodec='pcm_s16le', ac=1, ar='16k')
                .overwrite_output()
                .run(quiet=True)
            )
        
        pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization", use_auth_token=os.getenv("HUGGINGFACE_TOKEN"))
        diarization = pipeline(wav_file_path)
        speaker_segments = []
        for turn, _, speaker in diarization.itertracks(yield_label=True):
            speaker_segments.append({
                'start': turn.start,
                'end': turn.end,
                'speaker': speaker
            })
        
        # Remove the temporary WAV file if it was created
        if wav_file_path != audio_file_path:
            os.remove(wav_file_path)
            logging.info(f"Removed temporary WAV file: {wav_file_path}")
        
        return speaker_segments
    except Exception as e:
        logging.error(f"Error during speaker diarization: {e}")
        raise e

# Use AI to generate action items and follow-ups based on transcription
def generate_action_items_and_followups(transcription):
    logging.info("Generating action items and follow-ups using AI...")

    prompt = f"""
    Based on the following meeting transcription, extract all action items and follow-ups:
    
    {transcription}
    
    Provide a list of action items and follow-ups in the format:
    - Action Items: 
      * [Person] to [task], due by [date]
    - Follow-ups:
      * [Person] to follow up on [issue], due by [date]
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",  # Use GPT-4 for better understanding of context
            messages=[
                {"role": "system", "content": "You are an AI assistant that extracts action items and follow-ups from meeting transcriptions."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=500,
            temperature=0.2  # Low temperature for more deterministic output
        )
        
        if response.choices:
            ai_output = response.choices[0].message.content.strip()
            return ai_output
        else:
            raise ValueError("No valid AI output received.")
    except Exception as e:
        logging.error(f"Error generating action items and follow-ups: {e}")
        raise e

# New function to generate executive summary
def generate_executive_summary(transcription):
    logging.info("Generating executive summary...")
    prompt = f"""
    Based on the following meeting transcription, provide a brief executive summary of the key points discussed:
    
    {transcription}
    
    Limit the summary to 3-5 bullet points.
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an AI assistant that creates concise executive summaries."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=250,
            temperature=0.3
        )
        
        if response.choices:
            return response.choices[0].message.content.strip()
        else:
            raise ValueError("No valid AI output received.")
    except Exception as e:
        logging.error(f"Error generating executive summary: {e}")
        raise e

# New function to extract attendees
def extract_attendees(transcription):
    logging.info("Extracting attendees...")
    prompt = f"""
    Based on the following meeting transcription, extract the names and roles (if mentioned) of all attendees:
    
    {transcription}
    
    Provide the list in the format:
    - [Name] ([Role, if available])
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an AI assistant that extracts attendee information from meeting transcriptions."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=200,
            temperature=0.2
        )
        
        if response.choices:
            return response.choices[0].message.content.strip()
        else:
            raise ValueError("No valid AI output received.")
    except Exception as e:
        logging.error(f"Error extracting attendees: {e}")
        raise e

# Modify the generate_meeting_minutes function
def generate_meeting_minutes(transcription, audio_file_path):
    logging.info("Generating meeting minutes...")
    
    executive_summary = generate_executive_summary(transcription)
    attendees = extract_attendees(transcription)
    action_items_and_followups = generate_action_items_and_followups(transcription)
    
    return {
        'executive_summary': executive_summary,
        'attendees': attendees,
        'transcription': transcription,
        # 'speaker_diarization': speaker_segments,
        'action_items_and_followups': action_items_and_followups
    }

# Modify the save_as_docx function
def save_as_docx(minutes, audio_file_path):
    logging.info("Saving meeting minutes to Word document...")
    try:
        directory = os.path.dirname(audio_file_path)
        base_name = os.path.splitext(os.path.basename(audio_file_path))[0]
        output_path = os.path.join(directory, f"{base_name}_meeting_minutes.docx")

        doc = Document()

        # Add title
        doc.add_heading(f"Meeting Minutes: {base_name}", level=0)

        # Add executive summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(minutes['executive_summary'])

        # Add attendees
        doc.add_heading("Attendees", level=1)
        doc.add_paragraph(minutes['attendees'])

        # Add transcription
        doc.add_heading("Transcription", level=1)
        doc.add_paragraph(minutes['transcription'])

        # Add action items and follow-ups
        doc.add_heading("Action Items and Follow-ups", level=1)
        doc.add_paragraph(minutes['action_items_and_followups'])

        doc.save(output_path)
        logging.info(f"Meeting minutes saved as {output_path}")
    except Exception as e:
        logging.error(f"Error saving document: {e}")
        raise e

# Main function to process the audio file and generate meeting minutes
def process_meeting_audio_with_speakers(audio_file_path):
    transcription = transcribe_audio(audio_file_path)
    minutes = generate_meeting_minutes(transcription, audio_file_path)
    save_as_docx(minutes, audio_file_path)

# Function to handle command-line arguments
def main():
    parser = argparse.ArgumentParser(description="Transcribe an audio file and generate meeting minutes.")
    parser.add_argument('audio_file', type=str, help="Path to the audio file (e.g., .m4a, .wav, .mp3)")
    args = parser.parse_args()

    # Process the audio file
    process_meeting_audio_with_speakers(args.audio_file)

if __name__ == "__main__":
    main()